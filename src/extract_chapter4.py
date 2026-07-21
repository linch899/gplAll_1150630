import docx
import json
import re
import os
import sys

def replace_chinese_digits(s):
    mapping = {
        '零': '0', '〇': '0', '一': '1', '二': '2', '三': '3', 
        '四': '4', '五': '5', '六': '6', '七': '7', '八': '8', 
        '九': '9'
    }
    for k, v in mapping.items():
        s = s.replace(k, v)
    s = re.sub(r'([0-9])十([0-9])', r'\g<1>\g<2>', s)
    s = re.sub(r'([0-9])十', r'\g<1>0', s)
    s = re.sub(r'十([0-9])', r'1\g<1>', s)
    s = s.replace('十', '10')
    return s

def get_match_key(s):
    if not s:
        return ""
    s = re.sub(r'\s+', '', s)
    s = re.sub(r'[\(\)\-\[\]\{\}\uff08\uff09]', '', s)
    s = re.sub(r'(?:函|令)$', '', s)
    s = replace_chinese_digits(s)
    s = s.replace('第號', '第')
    return s

def clean_doc_id(doc_id):
    doc_id = re.sub(r'.*?日', '', doc_id)
    doc_id = doc_id.replace('第號', '第').strip()
    return doc_id

def find_match(last_line, gpl_map_exact, gpl_map_norm):
    raw_doc_id = clean_doc_id(last_line)
    std_doc_id = re.sub(r'\s+', '', raw_doc_id)
    
    if std_doc_id in gpl_map_exact:
        return gpl_map_exact[std_doc_id][0], "精確字號比對", std_doc_id
        
    for k, v in gpl_map_exact.items():
        if std_doc_id and (std_doc_id in k or k in std_doc_id):
            return v[0], "精確字號子字串比對", k
            
    norm_word = get_match_key(last_line)
    if norm_word in gpl_map_norm:
        return gpl_map_norm[norm_word][0], "歸一化精確比對", norm_word

    # Special fallbacks for docx typos / variants
    # 1. (88)工程企傳字第0018號 -> (88)工程企傳字第880018號
    if '工程企傳' in norm_word and '0018' in norm_word:
        for k, v in gpl_map_norm.items():
            if '工程企傳' in k and '880018' in k:
                return v[0], "企傳字號歸一化比對", k

    # 2. 09200305370 -> 09200305371
    if '0920030537' in norm_word:
        for k, v in gpl_map_norm.items():
            if '0920030537' in k:
                return v[0], "字號末位微差歸一化比對", k

    # 3. 0966002655910 -> 09600265910
    if '0966002655910' in norm_word or ('096' in norm_word and '265' in norm_word and '910' in norm_word):
        for k, v in gpl_map_norm.items():
            if '09600265910' in k:
                return v[0], "誤植字號修正比對", k

    # 4. 09700307700 -> 09700307740
    if '097003077' in norm_word:
        for k, v in gpl_map_norm.items():
            if '097003077' in k:
                return v[0], "字號末位微差歸一化比對", k
        
    matched_entries = []
    for db_key, db_entries in gpl_map_norm.items():
        if db_key and len(db_key) > 5 and (db_key in norm_word or norm_word in db_key):
            matched_entries.append((db_key, db_entries[0]))
            
    if matched_entries:
        matched_entries.sort(key=lambda x: len(x[0]), reverse=True)
        return matched_entries[0][1], "歸一化模糊比對", matched_entries[0][0]
        
    return None, "", ""

def get_date_val(item):
    date_str = item.get("發文日期", "").strip()
    if not date_str or len(date_str) < 4:
        return 0
    try:
        year_str = date_str[:-4]
        month_str = date_str[-4:-2]
        day_str = date_str[-2:]
        return int(year_str) * 10000 + int(month_str) * 100 + int(day_str)
    except ValueError:
        return 0

def extract_date(last_line):
    match = re.search(r'(\d+)年\s*(\d+)月\s*(\d+)日', last_line)
    if match:
        year = int(match.group(1))
        month = int(match.group(2))
        day = int(match.group(3))
        return f"{year}{month:02d}{day:02d}"
    return ""

def parse_leader_version(docx_path, gpl_map_exact, gpl_map_norm, comparison_log):
    doc = docx.Document(docx_path)
    
    sections = [
        "招標作業", "招標公告", "廠商領標與投標", "押標金、保證金",
        "疑義請求釋疑及處理結果", "訂定底價", "開標前之作業",
        "流標或廢標後重新招標", "其他招標階段事項"
    ]
    
    chinese_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九"]
    sec_names = {sec: f"第{chinese_nums[idx]}節 {sec}" for idx, sec in enumerate(sections)}
    
    current_sec = ""
    letters = []
    current_lines = []
    start_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
            
        pPr = p._p.get_or_add_pPr()
        numPr = pPr.find(docx.oxml.ns.qn('w:numPr'))
        ilvl_val = -1
        numId_val = -1
        if numPr is not None:
            ilvl = numPr.find(docx.oxml.ns.qn('w:ilvl'))
            numId = numPr.find(docx.oxml.ns.qn('w:numId'))
            ilvl_val = int(ilvl.get(docx.oxml.ns.qn('w:val'))) if ilvl is not None else -1
            numId_val = int(numId.get(docx.oxml.ns.qn('w:val'))) if numId is not None else -1
            
        if numId_val == 1 and ilvl_val == 0:
            continue
            
        if numId_val == 1 and ilvl_val == 1:
            if current_lines:
                letters.append((start_idx, current_lines, current_sec))
                current_lines = []
                start_idx = -1
            current_sec = sec_names.get(txt, f"第一節 {txt}")
            continue
            
        lines = [l.strip() for l in txt.split('\n') if l.strip()]
        if not lines:
            continue
            
        if not current_lines:
            start_idx = i
        current_lines.extend(lines)
        
        last_line = current_lines[-1]
        if re.search(r'(?:函|令)\s*\d+年\s*\d+月\s*\d+日?|字第\d+號', last_line) or re.search(r'\(?\d+\)?[^\s，、。]+字第\d+號', last_line):
            letters.append((start_idx, current_lines, current_sec))
            current_lines = []
            start_idx = -1
            
    if current_lines:
        letters.append((start_idx, current_lines, current_sec))
        
    extracted = []
    unmatched_count = 0
    
    for para_idx, lines, sec_name in letters:
        if not lines:
            continue
        last_line = lines[-1]
        
        match = re.search(r'(\(\d+\)[^\s，、。]+字第\d+號(?:函|令)?)', last_line)
        if not match:
            match = re.search(r'([^\s，、。]+字第\d+號(?:函|令)?)', last_line)
            
        has_valid_id = (match is not None)
        doc_id = match.group(1) if match else last_line
        doc_id = clean_doc_id(doc_id)
        
        db_match, method, matched_key = find_match(last_line, gpl_map_exact, gpl_map_norm)
        
        if db_match:
            new_entry = json.loads(json.dumps(db_match))
            new_entry["分類索引"] = {
                "章": "第四章 發包及廠商投標階段",
                "節": sec_name,
                "項": ""
            }
            extracted.append(new_entry)
            
            comparison_log.append({
                "來源文件": "第四章_組長版.docx",
                "段落索引": para_idx,
                "提取字號/最後一行": doc_id,
                "比對結果": "匹配成功",
                "資料庫發文字號": db_match.get("發文字號"),
                "比對方法": method,
                "匹配鍵值": matched_key,
                "預覽內容": lines[0][:40] + "..."
            })
        else:
            unmatched_count += 1
            placeholder_id = f"4-{unmatched_count:02d}"
            subject = lines[0]
            subject = re.sub(r'^主旨：', '', subject)
            content = "\n".join(lines)
            doc_date = extract_date(last_line)
            
            authority = "行政院公共工程委員會"
            for auth in ["內政部", "經濟部", "文化部", "法務部", "教育部", "原住民族委員會", "中華郵政股份有限公司", "行政院主計處"]:
                if auth in last_line:
                    authority = auth
                    break
                    
            new_entry = {
                "項次": placeholder_id,
                "分類索引": {
                    "章": "第四章 發包及廠商投標階段",
                    "節": sec_name,
                    "項": ""
                },
                "發文機關": authority,
                "發文字號": doc_id + ("函" if not doc_id.endswith(("函", "令")) else ""),
                "主題": subject,
                "依據採購法條文": "",
                "上網日期": "",
                "發文日期": doc_date,
                "連結網址": "",
                "內容": content,
                "廢止或補充之備註": ""
            }
            extracted.append(new_entry)
            
            is_actually_no_id = not has_valid_id and not re.search(r'字\s*第?\s*\d+\s*號', last_line)
            reason = "無發文字號" if is_actually_no_id else "比對不出(資料庫無此字號)"
            comparison_log.append({
                "來源文件": "第四章_組長版.docx",
                "段落索引": para_idx,
                "提取字號/最後一行": doc_id,
                "比對結果": f"未匹配 ({reason})",
                "資料庫發文字號": "",
                "比對方法": "",
                "匹配鍵值": "",
                "預覽內容": lines[0][:40] + "..."
            })
            
    return extracted, unmatched_count

def parse_new_version(docx_path, gpl_map_exact, gpl_map_norm, comparison_log):
    doc = docx.Document(docx_path)
    
    sections = [
        "招標作業", "招標公告", "廠商領標與投標", "押標金、保證金",
        "疑義請求釋疑及處理結果", "訂定底價", "開標前之作業",
        "流標或廢標後重新招標", "其他招標階段事項"
    ]
    chinese_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九"]
    sec_names = {sec: f"第{chinese_nums[idx]}節 {sec}" for idx, sec in enumerate(sections)}
    
    current_sec = ""
    letters = []
    current_lines = []
    start_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
            
        p_lines = [l.strip() for l in txt.split('\n') if l.strip()]
        if not p_lines:
            continue
            
        first = p_lines[0]
        if first.startswith("第四章") or first.startswith("章"):
            continue
            
        m_sec = re.match(r'^第[一二三四五六七八九十]+節\s*(.*)', first)
        if m_sec:
            if current_lines:
                letters.append((start_idx, current_lines, current_sec))
                current_lines = []
                start_idx = -1
            sec_core = m_sec.group(1).strip()
            current_sec = sec_names.get(sec_core, first)
            continue
            
        if first.startswith("主旨："):
            if current_lines:
                letters.append((start_idx, current_lines, current_sec))
                current_lines = []
                start_idx = -1
            start_idx = i
            
        if not current_lines and start_idx == -1:
            start_idx = i
        current_lines.extend(p_lines)
            
    if current_lines:
        letters.append((start_idx, current_lines, current_sec))
        
    extracted = []
    unmatched_count = 0
    
    for para_idx, lines, sec_name in letters:
        if not lines:
            continue
        
        # Find dispatch line inside lines
        dispatch_line = lines[-1]
        for l in reversed(lines):
            if re.search(r'(?:函|令)\s*\d+年\s*\d+月\s*\d+日?|字第\d+號', l) or re.search(r'\(?\d+\)?[^\s，、。]+字第\d+號', l):
                dispatch_line = l
                break
        
        match = re.search(r'(\(\d+\)[^\s，、。]+字第\d+號(?:函|令)?)', dispatch_line)
        if not match:
            match = re.search(r'([^\s，、。]+字第\d+號(?:函|令)?)', dispatch_line)
            
        has_valid_id = (match is not None)
        doc_id = match.group(1) if match else dispatch_line
        doc_id = clean_doc_id(doc_id)
        
        db_match, method, matched_key = find_match(dispatch_line, gpl_map_exact, gpl_map_norm)
        
        if db_match:
            new_entry = json.loads(json.dumps(db_match))
            new_entry["分類索引"] = {
                "章": "第四章 發包及廠商投標階段",
                "節": sec_name,
                "項": ""
            }
            extracted.append(new_entry)
            
            comparison_log.append({
                "來源文件": "第四章(112.1開始增加).docx",
                "段落索引": para_idx,
                "提取字號/最後一行": doc_id,
                "比對結果": "匹配成功",
                "資料庫發文字號": db_match.get("發文字號"),
                "比對方法": method,
                "匹配鍵值": matched_key,
                "預覽內容": "\n".join(lines)[:40] + "..."
            })
        else:
            unmatched_count += 1
            placeholder_id = f"4-new-{unmatched_count:02d}"
            content = "\n".join(lines)
            doc_date = extract_date(dispatch_line)
            subject = ""
            for line in lines:
                if line.startswith("主旨："):
                    subject = re.sub(r'^主旨：', '', line)
                    break
            if not subject:
                subject = lines[0]
                
            authority = "行政院公共工程委員會"
            for auth in ["內政部", "經濟部", "文化部", "法務部", "教育部", "原住民族委員會", "中華郵政股份有限公司", "行政院主計處"]:
                if auth in dispatch_line:
                    authority = auth
                    break
                    
            new_entry = {
                "項次": placeholder_id,
                "分類索引": {
                    "章": "第四章 發包及廠商投標階段",
                    "節": sec_name,
                    "項": ""
                },
                "發文機關": authority,
                "發文字號": doc_id + ("函" if not doc_id.endswith(("函", "令")) else ""),
                "主題": subject,
                "依據採購法條文": "",
                "上網日期": "",
                "發文日期": doc_date,
                "連結網址": "",
                "內容": content,
                "廢止或補充之備註": ""
            }
            extracted.append(new_entry)
            
            is_actually_no_id = not has_valid_id and not re.search(r'字\s*第?\s*\d+\s*號', dispatch_line)
            reason = "無發文字號" if is_actually_no_id else "比對不出(資料庫無此字號)"
            comparison_log.append({
                "來源文件": "第四章(112.1開始增加).docx",
                "段落索引": para_idx,
                "提取字號/最後一行": doc_id,
                "比對結果": f"未匹配 ({reason})",
                "資料庫發文字號": "",
                "比對方法": "",
                "匹配鍵值": "",
                "預覽內容": "\n".join(lines)[:40] + "..."
            })
            
    return extracted, unmatched_count

def main():
    docx_path_leader = r"D:\AI Workplace\antigravity\1150630_政府採購解釋函令前置作業\第四章\第四章_組長版.docx"
    docx_path_new = r"D:\AI Workplace\antigravity\1150630_政府採購解釋函令前置作業\第四章\第四章(112.1開始增加).docx"
    json_path = r"D:\AI Workplace\antigravity\1150630_政府採購解釋函令前置作業\gplAll_1150630.json"
    output_json_path = r"D:\AI Workplace\antigravity\1150630_政府採購解釋函令前置作業\第四章\第四章.json"
    output_log_path = r"D:\AI Workplace\antigravity\1150630_政府採購解釋函令前置作業\第四章\第四章比對紀錄.txt"

    print(f"Loading database from {json_path}...")
    with open(json_path, "r", encoding="utf-8") as f:
        gpl_data = json.load(f)

    gpl_map_exact = {}
    gpl_map_norm = {}
    for entry in gpl_data:
        doc_id = entry.get("發文字號", "").strip()
        if doc_id:
            std_id = re.sub(r'\s+', '', doc_id)
            gpl_map_exact.setdefault(std_id, []).append(entry)
            norm_id = get_match_key(doc_id)
            gpl_map_norm.setdefault(norm_id, []).append(entry)

    comparison_log = []

    print("Parsing fourth chapter leader version docx...")
    leader_entries, unmatched_leader = parse_leader_version(docx_path_leader, gpl_map_exact, gpl_map_norm, comparison_log)
    print(f"Leader Version completed. Extracted: {len(leader_entries)}, Unmatched: {unmatched_leader}")

    print("Parsing fourth chapter new 112.1 docx...")
    new_entries, unmatched_new = parse_new_version(docx_path_new, gpl_map_exact, gpl_map_norm, comparison_log)
    print(f"New Version completed. Extracted: {len(new_entries)}, Unmatched: {unmatched_new}")

    combined = leader_entries + new_entries
    print(f"Combined total entries: {len(combined)}")

    sections = [
        "招標作業", "招標公告", "廠商領標與投標", "押標金、保證金",
        "疑義請求釋疑及處理結果", "訂定底價", "開標前之作業",
        "流標或廢標後重新招標", "其他招標階段事項"
    ]
    chinese_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九"]
    section_order = [f"第{chinese_nums[idx]}節 {sec}" for idx, sec in enumerate(sections)]

    groups = {sec: [] for sec in section_order}
    other_groups = {}

    for item in combined:
        sec = item.get("分類索引", {}).get("節", "")
        if sec in groups:
            groups[sec].append(item)
        else:
            other_groups.setdefault(sec, []).append(item)

    sorted_combined = []
    print("Sorting each section internally by 民國 date...")
    for sec in section_order:
        sec_items = groups[sec]
        sec_items.sort(key=get_date_val)
        sorted_combined.extend(sec_items)
        print(f"  {sec}: {len(sec_items)} items sorted.")

    for sec, sec_items in other_groups.items():
        sec_items.sort(key=get_date_val)
        sorted_combined.extend(sec_items)
        print(f"  Other Section '{sec}': {len(sec_items)} items sorted.")

    print(f"Saving combined sorted Chapter 4 database to {output_json_path}...")
    os.makedirs(os.path.dirname(output_json_path), exist_ok=True)
    with open(output_json_path, "w", encoding="utf-8") as f:
        json.dump(sorted_combined, f, ensure_ascii=False, indent=2)

    print(f"Saving comparison log file to {output_log_path}...")
    leader_matched_count = sum(1 for item in comparison_log if item["來源文件"] == "第四章_組長版.docx" and item["比對結果"] == "匹配成功")
    new_matched_count = sum(1 for item in comparison_log if item["來源文件"] == "第四章(112.1開始增加).docx" and item["比對結果"] == "匹配成功")

    with open(output_log_path, "w", encoding="utf-8") as f:
        f.write("=========================================================================\n")
        f.write("                       第四章 解釋函令比對紀錄檔                         \n")
        f.write("=========================================================================\n\n")
        f.write(f"統計摘要：\n")
        f.write(f"  - 組長版：已匹配 {leader_matched_count} 筆，未匹配 {unmatched_leader} 筆\n")
        f.write(f"  - 新增版：已匹配 {new_matched_count} 筆，未匹配 {unmatched_new} 筆\n")
        f.write(f"  - 合計比對成功：{leader_matched_count + new_matched_count} 筆，合計未匹配：{unmatched_leader + unmatched_new} 筆\n\n")
        
        f.write("比對詳細清單：\n")
        f.write(f"{'來源文件':<25}{'段落':<6}{'比對結果':<15}{'提取/輸入字號':<35}{'資料庫發文字號':<30}{'比對方法':<20}{'預覽內容'}\n")
        f.write("-" * 160 + "\n")
        
        for item in comparison_log:
            src = item["來源文件"]
            p_idx = item["段落索引"]
            result = item["比對結果"]
            ext_id = item["提取字號/最後一行"]
            db_id = item["資料庫發文字號"]
            method = item["比對方法"]
            preview = item["預覽內容"].replace("\n", " ")
            
            f.write(f"{src:<25}{p_idx:<6}{result:<15}{ext_id:<35}{db_id:<30}{method:<20}{preview}\n")

    print("Chapter 4 Word to JSON processing completed successfully.")

if __name__ == "__main__":
    main()
