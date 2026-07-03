import json
import os
import re
import sys
import argparse
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Set UTF-8 encoding for stdout/stderr to avoid Windows terminal encoding errors
if sys.platform.startswith('win'):
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

# Regex for parsing document contents (using capturing groups to avoid variable-width lookbehinds)
HEADER_PATTERNS = {
    "主旨": re.compile(r"(?:^|\n)(主\s*旨\s*[:：])"),
    "說明": re.compile(r"(?:^|\n)(說\s*明\s*[:：])"),
    "辦法": re.compile(r"(?:^|\n)(辦\s*法\s*[:：])"),
    "正本": re.compile(r"(?:^|\n)(正\s*本\s*[:：])"),
    "副本": re.compile(r"(?:^|\n)(副\s*本\s*[:：])")
}

SIG_PATTERN = re.compile(r"(?:主任委員|院長|部長|署長)\s*[\u4e00-\u9fa5\s]{2,8}$")

MAP_DIGITS = {
    '0': '０', '1': '１', '2': '２', '3': '３', '4': '４',
    '5': '５', '6': '６', '7': '７', '8': '８', '9': '９'
}

def convert_digits(match):
    prefix = match.group(1) # either \n or empty string
    inner = match.group(2)
    converted_inner = "".join(MAP_DIGITS.get(c, c) for c in inner)
    return f"{prefix}（{converted_inner}）"

def convert_parentheses_and_digits(text):
    """
    Converts parentheses and digits to full-width format, but only if they are
    immediately preceded by a newline character (\n) or at the start of the string.
    """
    # 1. Convert (一), (二) to （一）, （二） after newline or start of string
    text = re.sub(r"(^|\n)\(([一二三四五六七八九十百]+)\)", r"\1（\2）", text)
    # 2. Convert (1), (2) to （１）, （２） after newline or start of string
    text = re.sub(r"(^|\n)\(([0-9０１２３４５６７８９]+)\)", convert_digits, text)
    return text

def set_run_font(run, font_name, size_pt=12, bold=False, color_rgb=None):
    """Sets Western and East Asian font name, size, bold and color on a run."""
    run.font.name = font_name
    if size_pt:
        run.font.size = Pt(size_pt)
    run.bold = bold
    if color_rgb:
        run.font.color.rgb = RGBColor.from_string(color_rgb)
    
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def clean_signature_and_copies(text):
    """Strips signature, positive, and negative copies from the end of the text."""
    text_strip = text.strip()
    
    # 1. Truncate from 正本/副本 onwards
    for key in ["正本", "副本"]:
        match = HEADER_PATTERNS[key].search(text_strip)
        if match:
            text_strip = text_strip[:match.start(1)].strip()
            
    # 2. Strip signature at the end if present
    match_sig = SIG_PATTERN.search(text_strip)
    if match_sig:
        text_strip = text_strip[:match_sig.start()].strip()
        
    return text_strip

def format_date_roc(date_str):
    """Formats a date string (YYMMDD or YYYMMDD) to ROC Chinese year format."""
    date_str = date_str.strip()
    if len(date_str) < 5:
        return date_str
    
    day = date_str[-2:]
    month = date_str[-4:-2]
    year = date_str[:-4]
    
    # Strip leading zero for month/day
    month_val = int(month)
    day_val = int(day)
    
    return f"{year}年{month_val:02d}月{day_val:02d}日"

def parse_letter_content(content):
    """
    Parses content field into a list of (segment_type, text) pairs.
    segment_type can be: '主旨', '說明', '辦法', or '內文'.
    """
    # Normalize newlines
    content = content.replace('\r\n', '\n').replace('\r', '\n')
    
    # Perform full-width conversion for line-start parentheses and numbers
    content = convert_parentheses_and_digits(content)
    
    cleaned = clean_signature_and_copies(content)
    
    # Find all header matches
    matches = []
    for key in ["主旨", "說明", "辦法"]:
        for match in HEADER_PATTERNS[key].finditer(cleaned):
            matches.append((match.start(1), match.end(1), key))
            
    # Sort matches by start position
    matches.sort(key=lambda x: x[0])
    
    segments = []
    if not matches:
        segments.append(('內文', cleaned))
        return segments
    
    # If there is text before the first match, extract it
    if matches[0][0] > 0:
        segments.append(('內文', cleaned[:matches[0][0]].strip()))
        
    for i in range(len(matches)):
        start_idx, end_idx, key = matches[i]
        next_start = matches[i+1][0] if i + 1 < len(matches) else len(cleaned)
        block_text = cleaned[end_idx:next_start].strip()
        segments.append((key, block_text))
        
    return segments

def split_segment_to_paragraphs(text):
    """Splits a segment text into paragraphs by newline and further by Level 1 markers."""
    lines = text.split('\n')
    paragraphs = []
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue
        # Split further ONLY on Level 1 markers (e.g. 一、, 二、)
        # to handle cases where list items are concatenated without newlines.
        chunks = re.split(r"(?=[一二三四五六七八九十百]+[、])", line_str)
        for chunk in chunks:
            chunk_str = chunk.strip()
            if chunk_str:
                paragraphs.append(chunk_str)
    return paragraphs

def add_heading_styled(doc, text, level):
    """Adds a Heading with Microsoft JhengHei, Bold, Red, and 24pt exact line spacing."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = Pt(24)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Pt(0)
    pf.first_line_indent = Pt(0)
    
    run = p.add_run(text)
    set_run_font(run, "微軟正黑體", size_pt=12, bold=True, color_rgb="FF0000")
    return p

def add_letter_header(doc, agency, date_str, word_num):
    """Adds a letter header paragraph in Microsoft JhengHei, Bold, Black, List Paragraph style."""
    p = doc.add_paragraph(style="List Paragraph")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Pt(36)
    pf.first_line_indent = Pt(-36)
    
    # Format ROC Date
    roc_date = format_date_roc(date_str)
    
    # Clean up word number (make sure it ends with 函 if needed)
    word_num = word_num.strip()
    if word_num and not word_num.endswith("函") and "字第" in word_num:
        word_num += "函"
        
    header_text = f"{agency}{roc_date}{word_num}"
    
    run = p.add_run(header_text)
    set_run_font(run, "微軟正黑體", size_pt=12, bold=True)
    return p

def get_list_level_and_indent(text):
    """
    Returns (level, left_indent_pt, first_line_indent_pt) based on the prefix of text.
    Alignments match '新版多層次排版範例.docx':
    Level 1 (一、): Left 36.0, FirstLine -24.0
    Level 2 (（一）): Left 48.2, FirstLine -34.0
    Level 3 (１、): Left 61.25, FirstLine -19.85
    Level 4 (（１）): Left 87.85, FirstLine -34.0
    Level 0 (Normal): Left 36.0, FirstLine -36.0
    """
    # Level 1: 一、, 二、
    if re.match(r"^[一二三四五六七八九十百]+、", text):
        return 1, 36.0, -24.0
        
    # Level 2: （一）, (一)
    if re.match(r"^[(（][一二三四五六七八九十百]+[)）]", text):
        matched = re.match(r"^[(（][一二三四五六七八九十百]+[)）]", text).group(0)
        inner = matched[1:-1].strip()
        roc_years = {
            "八八", "八九", "九十", "九一", "九二", "九三", "九四", "九五", "九六", "九七", "九八", "九九",
            "一〇〇", "一〇O", "一百", "一〇一", "一〇二", "一〇三", "一〇四", "一〇五", "一〇六", "一〇七", "一〇八", "一〇九",
            "一一〇", "一一一", "一一二", "一一三", "一一四", "一一五"
        }
        if inner not in roc_years:
            return 2, 48.2, -34.0
            
    # Level 3: １、, 1、, 1.
    if re.match(r"^[１２３４５６７８９０]+、", text):
        return 3, 61.25, -19.85
    if re.match(r"^\d+、", text):
        return 3, 61.25, -19.85
    if re.match(r"^\d+\.", text):
        return 3, 61.25, -19.85
        
    # Level 4: （１）, (1)
    if re.match(r"^[(（][１２３４５６７８９０\d]+[)）]", text):
        return 4, 87.85, -34.0
        
    # Not a list item
    return 0, 36.0, -36.0

def add_body_paragraph(doc, prefix, text, left_indent_pt=36.0, first_line_indent_pt=-36.0):
    """Adds a PMingLiU formatted body paragraph with specific indents."""
    p = doc.add_paragraph(style="List Paragraph")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    
    pf.left_indent = Pt(left_indent_pt)
    pf.first_line_indent = Pt(first_line_indent_pt)
        
    full_text = f"{prefix}{text}"
    run = p.add_run(full_text)
    set_run_font(run, "新細明體", size_pt=12)
    return p

def add_remarks(doc, remarks_text):
    """Adds formatting for the remarks section at the end of a letter."""
    remarks_text = remarks_text.strip()
    if not remarks_text:
        return
    
    # Split remarks by newlines
    lines = [line.strip() for line in remarks_text.split('\n') if line.strip()]
    if not lines:
        return
        
    if len(lines) == 1:
        # Single note
        p = doc.add_paragraph(style="List Paragraph")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.left_indent = Pt(36)
        pf.first_line_indent = Pt(-36)
        
        prefix = "備註："
        text = lines[0]
        if text.startswith("備註："):
            prefix = ""
        run = p.add_run(f"{prefix}{text}")
        set_run_font(run, "新細明體", size_pt=12)
    else:
        # Multiple notes
        for idx, line in enumerate(lines):
            p = doc.add_paragraph(style="List Paragraph")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.left_indent = Pt(48)
            
            if idx == 0:
                pf.first_line_indent = Pt(-48)
                prefix = "備註："
                text = line
                if text.startswith("備註："):
                    prefix = ""
                run = p.add_run(f"{prefix}{text}")
            else:
                # Determine alignment offset based on whether it starts with digit + .
                match_digit = re.match(r"^\d+\.", line)
                if match_digit:
                    pf.first_line_indent = Pt(-12)
                else:
                    pf.first_line_indent = Pt(0)
                run = p.add_run(line)
                
            set_run_font(run, "新細明體", size_pt=12)

def main():
    parser = argparse.ArgumentParser(description="Export PCC letters to formatted DOCX document.")
    parser.add_argument("--input", required=True, help="Path to input JSON file.")
    parser.add_argument("--output", required=True, help="Path to output DOCX file.")
    parser.add_argument("--range", help="Optional item index range, e.g. 1-100 (1-based index in JSON list).")
    args = parser.parse_args()
    
    # 1. Load JSON data
    if not os.path.exists(args.input):
        print(f"Error: Input JSON file '{args.input}' not found.")
        sys.exit(1)
        
    with open(args.input, "r", encoding="utf-8") as f:
        letters = json.load(f)
        
    print(f"Loaded {len(letters)} items from JSON.")
    
    # Apply range filter if specified
    if args.range:
        try:
            start_idx, end_idx = map(int, args.range.split('-'))
            letters = letters[start_idx-1 : end_idx]
            print(f"Filtered to index range {start_idx} to {end_idx} ({len(letters)} items).")
        except Exception as e:
            print(f"Error parsing range '{args.range}'. Use format 'start-end' (e.g. 1-10).")
            sys.exit(1)

    # 2. Initialize DOCX from template if it exists
    template_path = r"D:\AI Workplace\antigravity\1150630_政府採購解釋函令前置作業\第一章\新版多層次排版範例.docx"
    if os.path.exists(template_path):
        print(f"Loading styles and page margins from template: {template_path}")
        doc = docx.Document(template_path)
        # Clear body paragraphs while preserving styles and section setups
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
    else:
        print("Warning: Template not found. Creating a new document.")
        doc = docx.Document()
        # Set A4 margins
        section = doc.sections[0]
        section.page_width = docx.shared.Inches(8.27)
        section.page_height = docx.shared.Inches(11.69)
        section.top_margin = docx.shared.Inches(0.79)
        section.bottom_margin = docx.shared.Inches(0.98)
        section.left_margin = docx.shared.Inches(0.98)
        section.right_margin = docx.shared.Inches(0.98)

    # Keep track of current headings to prevent duplicates
    curr_chapter = None
    curr_section = None
    curr_subsection = None
    
    # Loop and write entries
    for idx, item in enumerate(letters):
        index_info = item.get("分類索引", {})
        chapter = index_info.get("章", "").strip()
        section = index_info.get("節", "").strip()
        subsection = index_info.get("項", "").strip()
        
        # Add Chapter Heading if changed
        if chapter and chapter != curr_chapter:
            add_heading_styled(doc, chapter, level=1)
            curr_chapter = chapter
            curr_section = None # Reset subheadings
            curr_subsection = None
            
        # Add Section Heading if changed
        if section and section != curr_section:
            add_heading_styled(doc, section, level=2)
            curr_section = section
            curr_subsection = None
            
        # Add Subsection Heading if changed
        if subsection and subsection != curr_subsection:
            add_heading_styled(doc, subsection, level=3)
            curr_subsection = subsection
            
        # Add an empty line before letter if it is not the very first paragraph
        if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text.strip():
            doc.add_paragraph()
            
        # Add Letter Header
        add_letter_header(
            doc, 
            item.get("發文機關", ""),
            item.get("發文日期", ""),
            item.get("發文字號", "")
        )
        
        # Parse and Add Letter Content segments
        segments = parse_letter_content(item.get("內容", ""))
        
        for seg_type, seg_text in segments:
            if not seg_text:
                continue
                
            if seg_type == '主旨':
                add_body_paragraph(doc, "主旨：", seg_text)
            elif seg_type in ['說明', '辦法']:
                # Split segment text by newlines and list items
                chunks = split_segment_to_paragraphs(seg_text)
                if not chunks:
                    continue
                    
                first_chunk = chunks[0]
                level, left, first = get_list_level_and_indent(first_chunk)
                
                if level > 0:
                    # No intro text, first paragraph is just the header (e.g. 說明：)
                    add_body_paragraph(doc, f"{seg_type}：", "", left_indent_pt=36.0, first_line_indent_pt=-36.0)
                    for chunk in chunks:
                        lvl, l_in, f_in = get_list_level_and_indent(chunk)
                        add_body_paragraph(doc, "", chunk, left_indent_pt=l_in, first_line_indent_pt=f_in)
                else:
                    # Has intro text, first chunk goes on same paragraph as header
                    add_body_paragraph(doc, f"{seg_type}：", first_chunk, left_indent_pt=36.0, first_line_indent_pt=-36.0)
                    for chunk in chunks[1:]:
                        lvl, l_in, f_in = get_list_level_and_indent(chunk)
                        add_body_paragraph(doc, "", chunk, left_indent_pt=l_in, first_line_indent_pt=f_in)
            else: # '內文'
                chunks = split_segment_to_paragraphs(seg_text)
                for chunk in chunks:
                    lvl, l_in, f_in = get_list_level_and_indent(chunk)
                    add_body_paragraph(doc, "", chunk, left_indent_pt=l_in, first_line_indent_pt=f_in)
                    
        # Add Remarks
        remarks_field = item.get("廢止或補充之備註", "")
        if remarks_field:
            add_remarks(doc, remarks_field)
            
    # Ensure any parent folders exist and save docx
    out_dir = os.path.dirname(args.output)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir)
        
    doc.save(args.output)
    print(f"Successfully generated DOCX file at: {args.output}")

if __name__ == "__main__":
    main()
